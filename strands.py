

import os
import json
import logging
import boto3
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
import asyncio
from dataclasses import dataclass

# Strands imports
from strands import Agent, tool
from strands.models import BedrockModel

# Additional imports for file handling
import PyPDF2
import docx
from bs4 import BeautifulSoup
import mammoth
import pandas as pd

# Configure logging
logging.getLogger("strands").setLevel(logging.DEBUG)
logging.basicConfig(
    format="%(levelname)s | %(name)s | %(message)s",
    handlers=[logging.StreamHandler()]
)

@dataclass
class TreatmentData:
    """Data structure for treatment information"""
    treatment_id: str
    source: str
    raw_content: str
    structured_content: str = ""
    risk_assessment: str = ""
    revenue_analysis: str = ""

@dataclass
class ProcessingResults:
    """Container for all processing results"""
    treatments: List[TreatmentData]
    final_report: str = ""
    approval_status: str = "pending"

class HealthcareAgentSystem:
    """Main system orchestrator for the healthcare agent hackathon"""
    
    def __init__(self):
        # Initialize AWS Bedrock session
        self.session = boto3.Session(
            aws_access_key_id='ASIAUFJ05VRONE JHNEP7',
            aws_secret_access_key='5FNHUatrDaqx97YKRAVPZ0a2gFgCUW3217RFN9MV',
            aws_session_token='IQ0Jb3JpZ2lux2VjEMn//////////wEaCXVZLWVhc3QtMSJGMEQCIG1TUFYEK5bxNkvX036wTbnjAEJCXOUF 1v2X450pBfJYAIA14t61Qitsis4jN1SBPIWNIWZ6NHn6',
            region_name='us-west-2'
        )
        
        # Create Bedrock model
        self.bedrock_model = BedrockModel(
            model_id="anthropic.claude-3-5-sonnet-20241022-v2:0",
            boto_session=self.session
        )
        
        # Initialize agents
        self.research_agent = None
        self.documenting_agent = None
        self.risk_assessment_agent = None
        self.revenue_identification_agent = None
        self.emailing_agent = None
        
        # Initialize data storage
        self.processing_results = ProcessingResults(treatments=[])
        
        # Setup output directory
        self.output_dir = Path("./hackathon_output")
        self.output_dir.mkdir(exist_ok=True)
        
        self._initialize_agents()
    
    def _initialize_agents(self):
        """Initialize all specialized agents"""
        self.research_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_file_reader_tool(),
                self._create_html_parser_tool(),
                self._create_pdf_parser_tool(),
                self._create_docx_parser_tool()
            ]
        )
        
        self.documenting_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_document_structuring_tool(),
                self._create_report_generator_tool()
            ]
        )
        
        self.risk_assessment_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_risk_analysis_tool(),
                self._create_medical_knowledge_tool()
            ]
        )
        
        self.revenue_identification_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_revenue_analysis_tool(),
                self._create_customer_segmentation_tool()
            ]
        )
        
        self.emailing_agent = Agent(
            model=self.bedrock_model,
            tools=[
                self._create_document_aggregation_tool(),
                self._create_word_export_tool()
            ]
        )

    # ==================== RESEARCH AGENT TOOLS ====================
    
    def _create_file_reader_tool(self):
        @tool
        def read_file(file_path: str) -> Dict[str, Any]:
            """
            Read various file formats from local filesystem.
            Supports .html, .pdf, .docx, .txt files.
            
            Args:
                file_path: Path to the file to read
                
            Returns:
                Dictionary containing file content and metadata
            """
            try:
                path = Path(file_path)
                if not path.exists():
                    return {"error": f"File not found: {file_path}"}
                
                file_extension = path.suffix.lower()
                content = ""
                
                if file_extension == '.html':
                    with open(path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif file_extension == '.txt':
                    with open(path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif file_extension == '.pdf':
                    with open(path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        content = ""
                        for page in pdf_reader.pages:
                            content += page.extract_text() + "\n"
                elif file_extension == '.docx':
                    doc = docx.Document(path)
                    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                else:
                    return {"error": f"Unsupported file format: {file_extension}"}
                
                return {
                    "file_path": str(path),
                    "file_type": file_extension,
                    "content": content,
                    "size": len(content),
                    "timestamp": datetime.now().isoformat()
                }
                
            except Exception as e:
                return {"error": f"Error reading file {file_path}: {str(e)}"}
        
        return read_file
    
    def _create_html_parser_tool(self):
        @tool
        def parse_html_content(html_content: str, target_treatments: List[str] = None) -> Dict[str, Any]:
            """
            Parse HTML content and extract treatment information.
            
            Args:
                html_content: Raw HTML content
                target_treatments: List of treatment IDs to look for
                
            Returns:
                Parsed treatment data
            """
            try:
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Extract text content
                text_content = soup.get_text(separator=' ', strip=True)
                
                # Look for treatment patterns
                treatments_found = []
                if target_treatments:
                    for treatment in target_treatments:
                        if treatment.lower() in text_content.lower():
                            treatments_found.append(treatment)
                
                # Extract structured data (tables, lists, etc.)
                tables = []
                for table in soup.find_all('table'):
                    table_data = []
                    for row in table.find_all('tr'):
                        row_data = [cell.get_text(strip=True) for cell in row.find_all(['td', 'th'])]
                        table_data.append(row_data)
                    tables.append(table_data)
                
                lists = []
                for ul in soup.find_all(['ul', 'ol']):
                    list_items = [li.get_text(strip=True) for li in ul.find_all('li')]
                    lists.append(list_items)
                
                return {
                    "text_content": text_content,
                    "treatments_found": treatments_found,
                    "tables": tables,
                    "lists": lists,
                    "title": soup.title.string if soup.title else "No title",
                    "headings": [h.get_text(strip=True) for h in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])]
                }
                
            except Exception as e:
                return {"error": f"Error parsing HTML: {str(e)}"}
        
        return parse_html_content
    
    def _create_pdf_parser_tool(self):
        @tool
        def extract_pdf_structured_data(pdf_content: str) -> Dict[str, Any]:
            """
            Extract structured information from PDF content.
            
            Args:
                pdf_content: Raw PDF text content
                
            Returns:
                Structured data extracted from PDF
            """
            try:
                # Split content into sections
                lines = pdf_content.split('\n')
                sections = []
                current_section = ""
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    # Check if line might be a heading (simple heuristic)
                    if len(line) < 100 and any(keyword in line.lower() for keyword in 
                                             ['treatment', 'procedure', 'therapy', 'diagnosis', 'symptom']):
                        if current_section:
                            sections.append(current_section)
                        current_section = line + "\n"
                    else:
                        current_section += line + "\n"
                
                if current_section:
                    sections.append(current_section)
                
                # Extract key medical terms
                medical_terms = []
                medical_keywords = ['diagnosis', 'treatment', 'therapy', 'procedure', 'medication', 
                                  'dosage', 'side effect', 'contraindication', 'efficacy']
                
                for line in lines:
                    for keyword in medical_keywords:
                        if keyword in line.lower():
                            medical_terms.append(line.strip())
                
                return {
                    "sections": sections,
                    "medical_terms": medical_terms,
                    "total_lines": len(lines),
                    "content_length": len(pdf_content)
                }
                
            except Exception as e:
                return {"error": f"Error processing PDF content: {str(e)}"}
        
        return extract_pdf_structured_data
    
    def _create_docx_parser_tool(self):
        @tool
        def parse_docx_content(docx_content: str) -> Dict[str, Any]:
            """
            Parse DOCX content for treatment information.
            
            Args:
                docx_content: Text content from DOCX file
                
            Returns:
                Parsed treatment data
            """
            try:
                paragraphs = [p.strip() for p in docx_content.split('\n') if p.strip()]
                
                # Identify potential treatment sections
                treatment_sections = []
                for i, paragraph in enumerate(paragraphs):
                    if any(keyword in paragraph.lower() for keyword in 
                          ['treatment', 'therapy', 'procedure', 'intervention']):
                        # Get context (previous and next paragraphs)
                        context = []
                        start = max(0, i-2)
                        end = min(len(paragraphs), i+3)
                        for j in range(start, end):
                            context.append(paragraphs[j])
                        treatment_sections.append({
                            'heading': paragraph,
                            'context': context
                        })
                
                return {
                    "paragraphs": paragraphs,
                    "treatment_sections": treatment_sections,
                    "paragraph_count": len(paragraphs)
                }
                
            except Exception as e:
                return {"error": f"Error parsing DOCX content: {str(e)}"}
        
        return parse_docx_content

    # ==================== DOCUMENTING AGENT TOOLS ====================
    
    def _create_document_structuring_tool(self):
        @tool
        def structure_treatment_data(raw_data: Dict[str, Any], treatment_id: str) -> Dict[str, Any]:
            """
            Structure raw treatment data into organized format.
            
            Args:
                raw_data: Raw extracted data from various sources
                treatment_id: ID of the treatment being processed
                
            Returns:
                Structured treatment document
            """
            try:
                structured_doc = {
                    "treatment_id": treatment_id,
                    "timestamp": datetime.now().isoformat(),
                    "data_sources": [],
                    "treatment_overview": "",
                    "clinical_details": {},
                    "regulatory_info": {},
                    "cost_information": {},
                    "effectiveness_data": {}
                }
                
                # Process content based on source type
                if isinstance(raw_data, dict):
                    if "content" in raw_data:
                        content = raw_data["content"]
                        
                        # Extract treatment overview
                        if "treatment" in content.lower():
                            lines = content.split('\n')
                            overview_lines = []
                            for line in lines:
                                if any(keyword in line.lower() for keyword in 
                                      ['treatment', 'therapy', 'procedure', 'overview']):
                                    overview_lines.append(line.strip())
                            structured_doc["treatment_overview"] = ' '.join(overview_lines[:3])
                        
                        # Extract clinical details
                        clinical_keywords = ['dosage', 'administration', 'duration', 'frequency']
                        for keyword in clinical_keywords:
                            if keyword in content.lower():
                                # Find sentences containing the keyword
                                sentences = content.split('.')
                                relevant_sentences = [s.strip() for s in sentences if keyword in s.lower()]
                                structured_doc["clinical_details"][keyword] = relevant_sentences[:2]
                        
                        # Extract effectiveness data
                        effectiveness_keywords = ['efficacy', 'success rate', 'outcome', 'result']
                        for keyword in effectiveness_keywords:
                            if keyword in content.lower():
                                sentences = content.split('.')
                                relevant_sentences = [s.strip() for s in sentences if keyword in s.lower()]
                                structured_doc["effectiveness_data"][keyword] = relevant_sentences[:2]
                        
                        structured_doc["data_sources"].append({
                            "file_path": raw_data.get("file_path", "unknown"),
                            "file_type": raw_data.get("file_type", "unknown"),
                            "processing_timestamp": datetime.now().isoformat()
                        })
                
                return structured_doc
                
            except Exception as e:
                return {"error": f"Error structuring treatment data: {str(e)}"}
        
        return structure_treatment_data
    
    def _create_report_generator_tool(self):
        @tool
        def generate_combined_report(structured_treatments: List[Dict[str, Any]]) -> str:
            """
            Generate a combined report from multiple structured treatments.
            
            Args:
                structured_treatments: List of structured treatment documents
                
            Returns:
                Combined report as formatted text
            """
            try:
                report_sections = []
                
                # Header
                report_sections.append("="*60)
                report_sections.append("HEALTHCARE TREATMENT ANALYSIS REPORT")
                report_sections.append("="*60)
                report_sections.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                report_sections.append(f"Total Treatments Analyzed: {len(structured_treatments)}")
                report_sections.append("")
                
                # Executive Summary
                report_sections.append("EXECUTIVE SUMMARY")
                report_sections.append("-" * 20)
                treatment_ids = [t.get("treatment_id", "Unknown") for t in structured_treatments]
                report_sections.append(f"Treatments Covered: {', '.join(treatment_ids)}")
                report_sections.append("")
                
                # Individual Treatment Analysis
                for i, treatment in enumerate(structured_treatments, 1):
                    report_sections.append(f"TREATMENT {i}: {treatment.get('treatment_id', 'Unknown')}")
                    report_sections.append("-" * 40)
                    
                    # Treatment Overview
                    overview = treatment.get("treatment_overview", "No overview available")
                    report_sections.append(f"Overview: {overview}")
                    report_sections.append("")
                    
                    # Clinical Details
                    clinical = treatment.get("clinical_details", {})
                    if clinical:
                        report_sections.append("Clinical Details:")
                        for key, value in clinical.items():
                            if isinstance(value, list):
                                report_sections.append(f"  - {key.title()}: {'; '.join(value)}")
                            else:
                                report_sections.append(f"  - {key.title()}: {value}")
                        report_sections.append("")
                    
                    # Effectiveness Data
                    effectiveness = treatment.get("effectiveness_data", {})
                    if effectiveness:
                        report_sections.append("Effectiveness Data:")
                        for key, value in effectiveness.items():
                            if isinstance(value, list):
                                report_sections.append(f"  - {key.title()}: {'; '.join(value)}")
                            else:
                                report_sections.append(f"  - {key.title()}: {value}")
                        report_sections.append("")
                    
                    # Data Sources
                    sources = treatment.get("data_sources", [])
                    if sources:
                        report_sections.append("Data Sources:")
                        for source in sources:
                            report_sections.append(f"  - {source.get('file_path', 'Unknown')}")
                        report_sections.append("")
                    
                    report_sections.append("")
                
                return "\n".join(report_sections)
                
            except Exception as e:
                return f"Error generating combined report: {str(e)}"
        
        return generate_combined_report

    # ==================== RISK ASSESSMENT AGENT TOOLS ====================
    
    def _create_risk_analysis_tool(self):
        @tool
        def analyze_treatment_risks(treatment_data: Dict[str, Any]) -> Dict[str, Any]:
            """
            Analyze risks associated with a treatment.
            
            Args:
                treatment_data: Structured treatment data
                
            Returns:
                Risk analysis results
            """
            try:
                risk_analysis = {
                    "treatment_id": treatment_data.get("treatment_id", "Unknown"),
                    "overall_risk_score": 0,  # Scale of 1-10
                    "medical_risks": [],
                    "financial_risks": [],
                    "regulatory_risks": [],
                    "recommendations": [],
                    "analysis_timestamp": datetime.now().isoformat()
                }
                
                # Analyze medical risks
                content = str(treatment_data)
                medical_risk_keywords = {
                    'high': ['contraindication', 'severe', 'fatal', 'death', 'emergency'],
                    'medium': ['side effect', 'adverse', 'caution', 'monitor', 'warning'],
                    'low': ['mild', 'temporary', 'reversible', 'minor']
                }
                
                medical_risk_score = 1
                for risk_level, keywords in medical_risk_keywords.items():
                    for keyword in keywords:
                        if keyword in content.lower():
                            risk_analysis["medical_risks"].append({
                                "keyword": keyword,
                                "level": risk_level,
                                "context": f"Found in treatment data: {keyword}"
                            })
                            if risk_level == 'high':
                                medical_risk_score += 3
                            elif risk_level == 'medium':
                                medical_risk_score += 2
                            elif risk_level == 'low':
                                medical_risk_score += 1
                
                # Analyze financial risks
                financial_keywords = ['expensive', 'cost', 'insurance', 'coverage', 'reimbursement']
                financial_risk_score = 1
                for keyword in financial_keywords:
                    if keyword in content.lower():
                        risk_analysis["financial_risks"].append({
                            "factor": keyword,
                            "impact": "medium",
                            "description": f"Financial consideration: {keyword}"
                        })
                        financial_risk_score += 1
                
                # Analyze regulatory risks
                regulatory_keywords = ['fda', 'approval', 'trial', 'experimental', 'investigational']
                regulatory_risk_score = 1
                for keyword in regulatory_keywords:
                    if keyword in content.lower():
                        risk_analysis["regulatory_risks"].append({
                            "factor": keyword,
                            "status": "requires_review",
                            "description": f"Regulatory consideration: {keyword}"
                        })
                        regulatory_risk_score += 1
                
                # Calculate overall risk score
                risk_analysis["overall_risk_score"] = min(10, 
                    (medical_risk_score + financial_risk_score + regulatory_risk_score) / 3)
                
                # Generate recommendations
                if risk_analysis["overall_risk_score"] >= 7:
                    risk_analysis["recommendations"].append("High risk treatment - requires extensive review")
                    risk_analysis["recommendations"].append("Consider alternative treatments")
                elif risk_analysis["overall_risk_score"] >= 4:
                    risk_analysis["recommendations"].append("Medium risk treatment - proceed with caution")
                    risk_analysis["recommendations"].append("Monitor patient outcomes closely")
                else:
                    risk_analysis["recommendations"].append("Low risk treatment - suitable for coverage")
                
                return risk_analysis
                
            except Exception as e:
                return {"error": f"Error analyzing treatment risks: {str(e)}"}
        
        return analyze_treatment_risks
    
    def _create_medical_knowledge_tool(self):
        @tool
        def consult_medical_knowledge(treatment_id: str, query: str) -> Dict[str, Any]:
            """
            Consult medical knowledge base for additional treatment information.
            
            Args:
                treatment_id: ID of the treatment
                query: Specific medical query
                
            Returns:
                Medical knowledge consultation results
            """
            try:
                # Simulated medical knowledge base
                medical_knowledge = {
                    "treatment_1": {
                        "category": "Cardiovascular",
                        "complexity": "High",
                        "typical_duration": "6-12 months",
                        "success_rate": "85%",
                        "common_complications": ["bleeding", "infection"]
                    },
                    "treatment_2": {
                        "category": "Orthopedic",
                        "complexity": "Medium",
                        "typical_duration": "3-6 months",
                        "success_rate": "90%",
                        "common_complications": ["swelling", "limited mobility"]
                    },
                    "treatment_3": {
                        "category": "Neurological",
                        "complexity": "High",
                        "typical_duration": "12-24 months",
                        "success_rate": "70%",
                        "common_complications": ["cognitive effects", "fatigue"]
                    },
                    "treatment_4": {
                        "category": "Oncological",
                        "complexity": "Very High",
                        "typical_duration": "6-18 months",
                        "success_rate": "75%",
                        "common_complications": ["nausea", "hair loss", "immunosuppression"]
                    },
                    "treatment_7": {
                        "category": "Endocrine",
                        "complexity": "Medium",
                        "typical_duration": "Lifelong",
                        "success_rate": "95%",
                        "common_complications": ["weight changes", "mood changes"]
                    },
                    "treatment_8": {
                        "category": "Respiratory",
                        "complexity": "Medium",
                        "typical_duration": "1-3 months",
                        "success_rate": "88%",
                        "common_complications": ["cough", "throat irritation"]
                    }
                }
                
                knowledge = medical_knowledge.get(treatment_id, {
                    "category": "General",
                    "complexity": "Unknown",
                    "typical_duration": "Variable",
                    "success_rate": "Unknown",
                    "common_complications": []
                })
                
                return {
                    "treatment_id": treatment_id,
                    "query": query,
                    "knowledge_base_info": knowledge,
                    "consultation_timestamp": datetime.now().isoformat(),
                    "confidence_level": "simulated_data"
                }
                
            except Exception as e:
                return {"error": f"Error consulting medical knowledge: {str(e)}"}
        
        return consult_medical_knowledge

    # ==================== REVENUE IDENTIFICATION AGENT TOOLS ====================
    
    def _create_revenue_analysis_tool(self):
        @tool
        def analyze_revenue_opportunities(treatment_data: Dict[str, Any], risk_data: Dict[str, Any]) -> Dict[str, Any]:
            """
            Analyze revenue opportunities for a treatment.
            
            Args:
                treatment_data: Structured treatment data
                risk_data: Risk analysis results
                
            Returns:
                Revenue opportunity analysis
            """
            try:
                revenue_analysis = {
                    "treatment_id": treatment_data.get("treatment_id", "Unknown"),
                    "market_potential": {},
                    "pricing_strategy": {},
                    "customer_segments": {},
                    "revenue_projections": {},
                    "business_case": {},
                    "analysis_timestamp": datetime.now().isoformat()
                }
                
                # Analyze market potential based on treatment complexity and success rate
                risk_score = risk_data.get("overall_risk_score", 5)
                
                if risk_score <= 3:
                    market_potential = "High"
                    base_price_multiplier = 1.2
                elif risk_score <= 6:
                    market_potential = "Medium"
                    base_price_multiplier = 1.0
                else:
                    market_potential = "Low"
                    base_price_multiplier = 0.8
                
                revenue_analysis["market_potential"] = {
                    "level": market_potential,
                    "risk_adjusted_score": 10 - risk_score,
                    "market_readiness": "ready" if risk_score <= 5 else "requires_evaluation"
                }
                
                # Pricing strategy
                base_coverage_amount = 10000  # Base amount in USD
                suggested_coverage = base_coverage_amount * base_price_multiplier
                
                revenue_analysis["pricing_strategy"] = {
                    "base_coverage_amount": base_coverage_amount,
                    "risk_multiplier": base_price_multiplier,
                    "suggested_coverage_limit": suggested_coverage,
                    "premium_adjustment": f"{int((base_price_multiplier - 1) * 100)}%"
                }
                
                # Customer segmentation
                customer_segments = {
                    "new_customers": {
                        "target_demographic": "Age 25-65, health-conscious",
                        "estimated_size": 50000 if market_potential == "High" else 25000,
                        "conversion_rate": "15%" if risk_score <= 4 else "8%"
                    },
                    "existing_customers": {
                        "upgrade_potential": "High" if risk_score <= 4 else "Medium",
                        "estimated_candidates": 20000 if market_potential == "High" else 10000,
                        "upsell_rate": "25%" if risk_score <= 4 else "15%"
                    }
                }
                
                revenue_analysis["customer_segments"] = customer_segments
                
                # Revenue projections
                new_customer_revenue = (customer_segments["new_customers"]["estimated_size"] * 
                                      (0.15 if risk_score <= 4 else 0.08) * suggested_coverage * 0.1)
                existing_customer_revenue = (customer_segments["existing_customers"]["estimated_candidates"] * 
                                           (0.25 if risk_score <= 4 else 0.15) * suggested_coverage * 0.05)
                
                total_projected_revenue = new_customer_revenue + existing_customer_revenue
                
                revenue_analysis["revenue_projections"] = {
                    "new_customer_revenue": new_customer_revenue,
                    "existing_customer_revenue": existing_customer_revenue,
                    "total_annual_projection": total_projected_revenue,
                    "roi_estimate": f"{int(total_projected_revenue / suggested_coverage * 100)}%"
                }
                
                # Business case
                business_case_strength = "Strong" if total_projected_revenue > 1000000 else "Moderate" if total_projected_revenue > 500000 else "Weak"
                
                revenue_analysis["business_case"] = {
                    "strength": business_case_strength,
                    "key_benefits": [
                        f"Market potential: {market_potential}",
                        f"Projected revenue: ${total_projected_revenue:,.2f}",
                        f"Risk level: {'Acceptable' if risk_score <= 5 else 'Requires mitigation'}"
                    ],
                    "recommendations": [
                        "Proceed with coverage" if business_case_strength == "Strong" else "Conduct further analysis",
                        "Target new customer acquisition" if new_customer_revenue > existing_customer_revenue else "Focus on existing customer upsell"
                    ]
                }
                
                return revenue_analysis
                
            except Exception as e:
                return {"error": f"Error analyzing revenue opportunities: {str(e)}"}
        
        return analyze_revenue_opportunities
    
    def _create_customer_segmentation_tool(self):
        @tool
        def segment_customers(treatment_analysis: Dict[str, Any]) -> Dict[str, Any]:
            """
            Perform detailed customer segmentation analysis.
            
            Args:
                treatment_analysis: Complete treatment analysis data
                
            Returns:
                Customer segmentation results
            """
            try:
                segmentation = {
                    "treatment_id": treatment_analysis.get("treatment_id", "Unknown"),
                    "segments": {},
                    "targeting_strategy": {},
                    "engagement_recommendations": {},
                    "segmentation_timestamp": datetime.now().isoformat()
                }
                
                # Define customer segments
                segments = {
                    "premium_seekers": {
                        "description": "Customers seeking comprehensive coverage",
                        "characteristics": ["High income", "Risk averse", "Values premium service"],
                        "size_estimate": "15%",
                        "targeting_priority": "High",
                        "messaging": "Comprehensive protection for peace of mind"
                    },
                    "cost_conscious": {
                        "description": "Price-sensitive customers",
                        "characteristics": ["Budget focused", "Compares options", "Values transparency"],
                        "size_estimate": "45%",
                        "targeting_priority": "Medium",
                        "messaging": "Affordable coverage without compromising quality"
                    },
                    "health_optimizers": {
                        "description": "Proactive health management customers",
                        "characteristics": ["Preventive care focused", "Tech-savvy", "Data-driven"],
                        "size_estimate": "25%",
                                                "targeting_priority": "High",
                        "messaging": "Advanced treatments for optimal health outcomes"
                    },
                    "chronic_condition": {
                        "description": "Patients with chronic conditions",
                        "characteristics": ["Long-term care needs", "Frequent healthcare utilization", "Medication dependent"],
                        "size_estimate": "15%",
                        "targeting_priority": "Medium",
                        "messaging": "Specialized coverage for your ongoing needs"
                    }
                }
                
                segmentation["segments"] = segments
                
                # Determine targeting strategy based on treatment characteristics
                risk_score = treatment_analysis.get("risk_data", {}).get("overall_risk_score", 5)
                market_potential = treatment_analysis.get("revenue_data", {}).get("market_potential", {}).get("level", "Medium")
                
                if risk_score <= 3 and market_potential == "High":
                    targeting_strategy = {
                        "primary_segment": "premium_seekers",
                        "secondary_segment": "health_optimizers",
                        "approach": "Premium positioning with value-added services"
                    }
                elif risk_score <= 5:
                    targeting_strategy = {
                        "primary_segment": "health_optimizers",
                        "secondary_segment": "cost_conscious",
                        "approach": "Balanced value proposition"
                    }
                else:
                    targeting_strategy = {
                        "primary_segment": "chronic_condition",
                        "secondary_segment": "premium_seekers",
                        "approach": "Niche positioning with risk mitigation"
                    }
                
                segmentation["targeting_strategy"] = targeting_strategy
                
                # Engagement recommendations
                engagement_recommendations = []
                
                if targeting_strategy["primary_segment"] == "premium_seekers":
                    engagement_recommendations.extend([
                        "Develop premium service package",
                        "Offer concierge service options",
                        "Highlight exclusivity and comprehensive coverage"
                    ])
                elif targeting_strategy["primary_segment"] == "health_optimizers":
                    engagement_recommendations.extend([
                        "Create educational content about treatment benefits",
                        "Develop digital health tracking tools",
                        "Offer personalized treatment plans"
                    ])
                else:
                    engagement_recommendations.extend([
                        "Focus on cost-effectiveness",
                        "Provide transparent pricing information",
                        "Offer flexible payment options"
                    ])
                
                segmentation["engagement_recommendations"] = {
                    "primary_segment": engagement_recommendations,
                    "secondary_segment": [
                        "Tailored messaging for secondary audience",
                        "Cross-promotion with related products"
                    ]
                }
                
                return segmentation
                
            except Exception as e:
                return {"error": f"Error performing customer segmentation: {str(e)}"}
        
        return segment_customers

    # ==================== EMAILING AGENT TOOLS ====================
    
    def _create_document_aggregation_tool(self):
        @tool
        def aggregate_documents(treatment_reports: List[Dict[str, Any]], risk_assessments: List[Dict[str, Any]], 
                              revenue_analyses: List[Dict[str, Any]]) -> Dict[str, Any]:
            """
            Aggregate all documents into a single package for approval.
            
            Args:
                treatment_reports: List of treatment reports
                risk_assessments: List of risk assessments
                revenue_analyses: List of revenue analyses
                
            Returns:
                Aggregated document package
            """
            try:
                aggregated = {
                    "metadata": {
                        "generated_at": datetime.now().isoformat(),
                        "total_treatments": len(treatment_reports),
                        "total_pages": 0,
                        "document_version": "1.0"
                    },
                    "treatments": [],
                    "summary": {}
                }
                
                # Process each treatment
                for i in range(len(treatment_reports)):
                    treatment_data = {
                        "treatment_id": treatment_reports[i].get("treatment_id", f"Unknown_{i}"),
                        "report": treatment_reports[i],
                        "risk_assessment": risk_assessments[i] if i < len(risk_assessments) else {},
                        "revenue_analysis": revenue_analyses[i] if i < len(revenue_analyses) else {}
                    }
                    aggregated["treatments"].append(treatment_data)
                
                # Create executive summary
                total_risk_score = sum(
                    r.get("overall_risk_score", 5) for r in risk_assessments) / max(1, len(risk_assessments))
                total_revenue = sum(
                    r.get("revenue_projections", {}).get("total_annual_projection", 0) 
                    for r in revenue_analyses)
                
                aggregated["summary"] = {
                    "average_risk_score": total_risk_score,
                    "total_projected_revenue": total_revenue,
                    "recommendation": "APPROVE" if total_risk_score <= 5 else "REVIEW",
                    "key_treatments": [t["treatment_id"] for t in aggregated["treatments"]][:3],
                    "generated_by": "HealthcareAgentSystem"
                }
                
                return aggregated
                
            except Exception as e:
                return {"error": f"Error aggregating documents: {str(e)}"}
        
        return aggregate_documents
    
    def _create_word_export_tool(self):
        @tool
        def generate_word_document(aggregated_data: Dict[str, Any], output_path: str) -> Dict[str, Any]:
            """
            Generate a Word document from aggregated data.
            
            Args:
                aggregated_data: Aggregated document package
                output_path: Path to save the Word document
                
            Returns:
                Generation results
            """
            try:
                doc = docx.Document()
                
                # Add title
                doc.add_heading('Healthcare Treatment Analysis Package', 0)
                
                # Add metadata
                doc.add_heading('Document Information', 1)
                metadata = aggregated_data.get("metadata", {})
                doc.add_paragraph(f"Generated at: {metadata.get('generated_at', 'Unknown')}")
                doc.add_paragraph(f"Total treatments analyzed: {metadata.get('total_treatments', 0)}")
                doc.add_paragraph(f"Document version: {metadata.get('document_version', '1.0')}")
                
                # Add executive summary
                doc.add_heading('Executive Summary', 1)
                summary = aggregated_data.get("summary", {})
                doc.add_paragraph(f"Average risk score: {summary.get('average_risk_score', 0):.1f}/10")
                doc.add_paragraph(f"Total projected annual revenue: ${summary.get('total_projected_revenue', 0):,.2f}")
                doc.add_paragraph(f"Recommendation: {summary.get('recommendation', 'REVIEW')}")
                
                # Add treatment details
                doc.add_heading('Treatment Details', 1)
                for treatment in aggregated_data.get("treatments", [])[:5]:  # Limit to first 5 for demo
                    doc.add_heading(treatment.get("treatment_id", "Unknown Treatment"), 2)
                    
                    # Add report highlights
                    report = treatment.get("report", {})
                    doc.add_paragraph("Treatment Overview:", style='Heading 3')
                    doc.add_paragraph(report.get("treatment_overview", "No overview available"))
                    
                    # Add risk highlights
                    risk = treatment.get("risk_assessment", {})
                    doc.add_paragraph("Risk Assessment:", style='Heading 3')
                    doc.add_paragraph(f"Overall risk score: {risk.get('overall_risk_score', 0)}")
                    doc.add_paragraph("Key risks:")
                    for r in risk.get("medical_risks", [])[:3]:
                        doc.add_paragraph(f"- {r.get('keyword', '')} ({r.get('level', '')})", style='List Bullet')
                    
                    # Add revenue highlights
                    revenue = treatment.get("revenue_analysis", {})
                    doc.add_paragraph("Revenue Potential:", style='Heading 3')
                    doc.add_paragraph(f"Market potential: {revenue.get('market_potential', {}).get('level', 'Unknown')}")
                    doc.add_paragraph(f"Projected revenue: ${revenue.get('revenue_projections', {}).get('total_annual_projection', 0):,.2f}")
                
                # Save document
                doc.save(output_path)
                
                return {
                    "output_path": output_path,
                    "file_size": os.path.getsize(output_path),
                    "treatments_included": len(aggregated_data.get("treatments", [])),
                    "generation_status": "success"
                }
                
            except Exception as e:
                return {"error": f"Error generating Word document: {str(e)}"}
        
        return generate_word_document

    # ==================== MAIN PROCESSING METHODS ====================
    
    async def process_treatments(self, treatment_ids: List[str], source_files: List[str]):
        """
        Main processing pipeline for treatments.
        
        Args:
            treatment_ids: List of treatment IDs to process
            source_files: List of source file paths
        """
        try:
            # Step 1: Research Agent - Extract raw data
            raw_treatment_data = await self._extract_treatment_data(source_files, treatment_ids)
            
            # Step 2: Documenting Agent - Structure data
            structured_treatments = await self._structure_treatment_data(raw_treatment_data, treatment_ids)
            
            # Step 3: Risk Assessment Agent - Analyze risks
            risk_assessments = await self._assess_treatment_risks(structured_treatments)
            
            # Step 4: Revenue Identification Agent - Analyze revenue
            revenue_analyses = await self._analyze_revenue_opportunities(structured_treatments, risk_assessments)
            
            # Step 5: Generate final report
            final_report = await self._generate_final_report(structured_treatments, risk_assessments, revenue_analyses)
            
            # Step 6: Prepare approval package
            approval_package = await self._prepare_approval_package(structured_treatments, risk_assessments, revenue_analyses)
            
            # Step 7: Email results
            await self._email_results(approval_package)
            
            return {
                "status": "success",
                "treatments_processed": len(treatment_ids),
                "final_report_path": str(self.output_dir / "final_report.docx"),
                "approval_package": approval_package
            }
            
        except Exception as e:
            logging.error(f"Error in processing pipeline: {str(e)}")
            return {"status": "error", "error": str(e)}
    
    async def _extract_treatment_data(self, source_files: List[str], treatment_ids: List[str]) -> List[Dict[str, Any]]:
        """Extract raw treatment data from source files"""
        extracted_data = []
        
        for file_path in source_files:
            # Read file content
            file_content = await self.research_agent.run(
                "read_file",
                file_path=file_path
            )
            
            if "error" in file_content:
                logging.warning(f"Skipping file {file_path}: {file_content['error']}")
                continue
            
            # Parse content based on file type
            if file_content["file_type"] == ".html":
                parsed_data = await self.research_agent.run(
                    "parse_html_content",
                    html_content=file_content["content"],
                    target_treatments=treatment_ids
                )
            elif file_content["file_type"] == ".pdf":
                parsed_data = await self.research_agent.run(
                    "extract_pdf_structured_data",
                    pdf_content=file_content["content"]
                )
            elif file_content["file_type"] == ".docx":
                parsed_data = await self.research_agent.run(
                    "parse_docx_content",
                    docx_content=file_content["content"]
                )
            else:
                parsed_data = {"content": file_content["content"]}
            
            if "error" in parsed_data:
                logging.warning(f"Error parsing {file_path}: {parsed_data['error']}")
                continue
            
            extracted_data.append({
                "file_path": file_path,
                "file_type": file_content["file_type"],
                "content": parsed_data,
                "treatments_found": parsed_data.get("treatments_found", [])
            })
        
        return extracted_data
    
    async def _structure_treatment_data(self, raw_data: List[Dict[str, Any]], treatment_ids: List[str]) -> List[Dict[str, Any]]:
        """Structure raw treatment data into organized format"""
        structured_treatments = []
        
        for treatment_id in treatment_ids:
            # Find all raw data mentioning this treatment
            treatment_raw_data = [d for d in raw_data if treatment_id in d.get("treatments_found", [])]
            
            if not treatment_raw_data:
                logging.warning(f"No data found for treatment {treatment_id}")
                continue
            
            # Structure each piece of raw data
            structured_pieces = []
            for raw_piece in treatment_raw_data:
                structured = await self.documenting_agent.run(
                    "structure_treatment_data",
                    raw_data=raw_piece,
                    treatment_id=treatment_id
                )
                
                if "error" not in structured:
                    structured_pieces.append(structured)
            
            if structured_pieces:
                # Combine structured pieces
                combined_treatment = {
                    "treatment_id": treatment_id,
                    "sources": [s["data_sources"] for s in structured_pieces],
                    "treatment_overview": "\n".join([s.get("treatment_overview", "") for s in structured_pieces]),
                    "clinical_details": {},
                    "effectiveness_data": {}
                }
                
                # Merge clinical details
                clinical_keys = set()
                for piece in structured_pieces:
                    clinical_keys.update(piece.get("clinical_details", {}).keys())
                
                for key in clinical_keys:
                    combined_values = []
                    for piece in structured_pieces:
                        combined_values.extend(piece.get("clinical_details", {}).get(key, []))
                    combined_treatment["clinical_details"][key] = combined_values[:5]  # Limit to top 5
                
                # Merge effectiveness data
                effectiveness_keys = set()
                for piece in structured_pieces:
                    effectiveness_keys.update(piece.get("effectiveness_data", {}).keys())
                
                for key in effectiveness_keys:
                    combined_values = []
                    for piece in structured_pieces:
                        combined_values.extend(piece.get("effectiveness_data", {}).get(key, []))
                    combined_treatment["effectiveness_data"][key] = combined_values[:5]
                
                structured_treatments.append(combined_treatment)
        
        return structured_treatments
    
    async def _assess_treatment_risks(self, structured_treatments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Assess risks for each structured treatment"""
        risk_assessments = []
        
        for treatment in structured_treatments:
            # Basic risk analysis
            risk_analysis = await self.risk_assessment_agent.run(
                "analyze_treatment_risks",
                treatment_data=treatment
            )
            
            if "error" in risk_analysis:
                logging.warning(f"Error assessing risks for {treatment['treatment_id']}: {risk_analysis['error']}")
                continue
            
            # Consult medical knowledge base
            medical_knowledge = await self.risk_assessment_agent.run(
                "consult_medical_knowledge",
                treatment_id=treatment["treatment_id"],
                query="Provide standard medical context for this treatment"
            )
            
            if "error" not in medical_knowledge:
                risk_analysis["medical_context"] = medical_knowledge["knowledge_base_info"]
            
            risk_assessments.append(risk_analysis)
        
        return risk_assessments
    
    async def _analyze_revenue_opportunities(self, structured_treatments: List[Dict[str, Any]], 
                                           risk_assessments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Analyze revenue opportunities for each treatment"""
        revenue_analyses = []
        
        for i, treatment in enumerate(structured_treatments):
            if i >= len(risk_assessments):
                break
                
            # Basic revenue analysis
            revenue_analysis = await self.revenue_identification_agent.run(
                "analyze_revenue_opportunities",
                treatment_data=treatment,
                risk_data=risk_assessments[i]
            )
            
            if "error" in revenue_analysis:
                logging.warning(f"Error analyzing revenue for {treatment['treatment_id']}: {revenue_analysis['error']}")
                continue
            
            # Customer segmentation
            segmentation = await self.revenue_identification_agent.run(
                "segment_customers",
                treatment_analysis={
                    "treatment_id": treatment["treatment_id"],
                    "risk_data": risk_assessments[i],
                    "revenue_data": revenue_analysis
                }
            )
            
            if "error" not in segmentation:
                revenue_analysis["customer_segmentation"] = segmentation
            
            revenue_analyses.append(revenue_analysis)
        
        return revenue_analyses
    
    async def _generate_final_report(self, structured_treatments: List[Dict[str, Any]], 
                                   risk_assessments: List[Dict[str, Any]], 
                                   revenue_analyses: List[Dict[str, Any]]) -> str:
        """Generate final combined report"""
        combined_data = []
        
        for i, treatment in enumerate(structured_treatments):
            combined = {
                "treatment_id": treatment["treatment_id"],
                "report": treatment,
                "risk_assessment": risk_assessments[i] if i < len(risk_assessments) else {},
                "revenue_analysis": revenue_analyses[i] if i < len(revenue_analyses) else {}
            }
            combined_data.append(combined)
        
        report = await self.documenting_agent.run(
            "generate_combined_report",
            structured_treatments=combined_data
        )
        
        # Save report to file
        report_path = self.output_dir / "final_report.txt"
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(report)
        
        return str(report_path)
    
    async def _prepare_approval_package(self, structured_treatments: List[Dict[str, Any]], 
                                      risk_assessments: List[Dict[str, Any]], 
                                      revenue_analyses: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Prepare complete approval package"""
        approval_package = await self.emailing_agent.run(
            "aggregate_documents",
            treatment_reports=structured_treatments,
            risk_assessments=risk_assessments,
            revenue_analyses=revenue_analyses
        )
        
        if "error" in approval_package:
            raise Exception(f"Error aggregating documents: {approval_package['error']}")
        
        # Generate Word document
        doc_path = self.output_dir / "approval_package.docx"
        doc_result = await self.emailing_agent.run(
            "generate_word_document",
            aggregated_data=approval_package,
            output_path=str(doc_path)
        )
        
        if "error" in doc_result:
            raise Exception(f"Error generating Word document: {doc_result['error']}")
        
        approval_package["document_path"] = str(doc_path)
        return approval_package
    
    async def _email_results(self, approval_package: Dict[str, Any]):
        """Email results to business stakeholders"""
        # In a real implementation, this would use AWS SES or similar
        # For demo purposes, we'll just log the would-be email
        
        email_content = f"""
        Healthcare Treatment Analysis Results
        
        Treatments Processed: {approval_package.get('metadata', {}).get('total_treatments', 0)}
        Average Risk Score: {approval_package.get('summary', {}).get('average_risk_score', 0):.1f}/10
        Total Projected Revenue: ${approval_package.get('summary', {}).get('total_projected_revenue', 0):,.2f}
        Recommendation: {approval_package.get('summary', {}).get('recommendation', 'REVIEW')}
        
        The complete approval package is attached.
        """
        
        logging.info(f"Would send email with content:\n{email_content}")
        logging.info(f"With attachment at: {approval_package.get('document_path', '')}")

# ==================== MAIN EXECUTION ====================

async def main():
    """Main execution function"""
    try:
        # Initialize the system
        system = HealthcareAgentSystem()
        
        # Define inputs
        treatment_ids = ["treatment_2", "treatment_4", "treatment_7"]
        source_files = [
            "Government Portal.html",
            "Competitor Portal.html"
        ]
        
        # Process treatments
        results = await system.process_treatments(treatment_ids, source_files)
        
        if results["status"] == "success":
            print(f"Successfully processed {results['treatments_processed']} treatments")
            print(f"Final report generated at: {results['final_report_path']}")
        else:
            print(f"Processing failed: {results.get('error', 'Unknown error')}")
            
    except Exception as e:
        print(f"Fatal error: {str(e)}")

if __name__ == "__main__":
    asyncio.run(main())